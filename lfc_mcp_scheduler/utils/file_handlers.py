"""File import/export handlers for different formats."""

import json
import csv
from datetime import datetime, date
from typing import List, Dict, Any, Optional
from pathlib import Path
import pandas as pd
from openpyxl import Workbook, load_workbook
from docx import Document
from docx.shared import Inches

from ..models.schedule import Schedule, ScheduleEvent, Priority, EventStatus, RepeatType


class ExcelHandler:
    """Handle Excel file import/export."""
    
    @staticmethod
    def export_to_excel(schedule: Schedule, file_path: str) -> None:
        """Export schedule to Excel file."""
        wb = Workbook()
        ws = wb.active
        ws.title = "Schedule"
        
        # Headers
        headers = [
            "ID", "标题", "描述", "开始时间", "结束时间", "地点", 
            "优先级", "状态", "标签", "重复类型", "重复结束日期",
            "参与者", "备注", "创建时间", "更新时间"
        ]
        ws.append(headers)
        
        # Data rows
        for event in schedule.events:
            row = [
                event.id,
                event.title,
                event.description or "",
                event.start_datetime.strftime("%Y-%m-%d %H:%M:%S"),
                event.end_datetime.strftime("%Y-%m-%d %H:%M:%S") if event.end_datetime else "",
                event.location or "",
                event.priority.value,
                event.status.value,
                ", ".join(event.tags),
                event.repeat_type.value,
                event.repeat_until.strftime("%Y-%m-%d") if event.repeat_until else "",
                ", ".join(event.attendees),
                event.notes or "",
                event.created_at.strftime("%Y-%m-%d %H:%M:%S"),
                event.updated_at.strftime("%Y-%m-%d %H:%M:%S")
            ]
            ws.append(row)
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        wb.save(file_path)
    
    @staticmethod
    def import_from_excel(file_path: str) -> Schedule:
        """Import schedule from Excel file."""
        wb = load_workbook(file_path)
        ws = wb.active
        
        schedule = Schedule(name=f"Imported from {Path(file_path).name}")
        
        # Skip header row
        rows = list(ws.iter_rows(min_row=2, values_only=True))
        
        for row in rows:
            if not row[0]:  # Skip empty rows
                continue
            
            try:
                event = ScheduleEvent(
                    id=str(row[0]) if row[0] else None,
                    title=str(row[1]) if row[1] else "未命名事件",
                    description=str(row[2]) if row[2] else None,
                    start_datetime=datetime.strptime(str(row[3]), "%Y-%m-%d %H:%M:%S") if row[3] else datetime.now(),
                    end_datetime=datetime.strptime(str(row[4]), "%Y-%m-%d %H:%M:%S") if row[4] else None,
                    location=str(row[5]) if row[5] else None,
                    priority=Priority(str(row[6])) if row[6] and str(row[6]) in [p.value for p in Priority] else Priority.MEDIUM,
                    status=EventStatus(str(row[7])) if row[7] and str(row[7]) in [s.value for s in EventStatus] else EventStatus.PENDING,
                    tags=str(row[8]).split(", ") if row[8] else [],
                    repeat_type=RepeatType(str(row[9])) if row[9] and str(row[9]) in [r.value for r in RepeatType] else RepeatType.NONE,
                    repeat_until=datetime.strptime(str(row[10]), "%Y-%m-%d").date() if row[10] else None,
                    attendees=str(row[11]).split(", ") if row[11] else [],
                    notes=str(row[12]) if row[12] else None
                )
                schedule.add_event(event)
            except Exception as e:
                print(f"Error importing row {row}: {e}")
                continue
        
        return schedule


class TxtHandler:
    """Handle TXT file import/export."""
    
    @staticmethod
    def export_to_txt(schedule: Schedule, file_path: str) -> None:
        """Export schedule to TXT file."""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(f"日程安排: {schedule.name}\n")
            f.write(f"描述: {schedule.description or '无'}\n")
            f.write(f"创建时间: {schedule.created_at.strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("=" * 50 + "\n\n")
            
            for i, event in enumerate(schedule.events, 1):
                f.write(f"{i}. {event.title}\n")
                f.write(f"   时间: {event.start_datetime.strftime('%Y-%m-%d %H:%M')}")
                if event.end_datetime:
                    f.write(f" - {event.end_datetime.strftime('%Y-%m-%d %H:%M')}")
                f.write("\n")
                
                if event.description:
                    f.write(f"   描述: {event.description}\n")
                if event.location:
                    f.write(f"   地点: {event.location}\n")
                f.write(f"   优先级: {event.priority.value}\n")
                f.write(f"   状态: {event.status.value}\n")
                
                if event.tags:
                    f.write(f"   标签: {', '.join(event.tags)}\n")
                if event.attendees:
                    f.write(f"   参与者: {', '.join(event.attendees)}\n")
                if event.notes:
                    f.write(f"   备注: {event.notes}\n")
                
                f.write("\n" + "-" * 30 + "\n\n")
    
    @staticmethod
    def import_from_txt(file_path: str) -> Schedule:
        """Import schedule from TXT file (simple format)."""
        schedule = Schedule(name=f"Imported from {Path(file_path).name}")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Simple parsing - each line starting with number is an event
        lines = content.split('\n')
        current_event = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line starts with a number (new event)
            if line[0].isdigit() and '. ' in line:
                if current_event:
                    schedule.add_event(current_event)
                
                title = line.split('. ', 1)[1]
                current_event = ScheduleEvent(
                    title=title,
                    start_datetime=datetime.now()
                )
            elif current_event and line.startswith('   时间:'):
                # Parse time information
                time_str = line.replace('   时间:', '').strip()
                try:
                    if ' - ' in time_str:
                        start_str, end_str = time_str.split(' - ')
                        current_event.start_datetime = datetime.strptime(start_str, '%Y-%m-%d %H:%M')
                        current_event.end_datetime = datetime.strptime(end_str, '%Y-%m-%d %H:%M')
                    else:
                        current_event.start_datetime = datetime.strptime(time_str, '%Y-%m-%d %H:%M')
                except:
                    pass
            elif current_event:
                # Parse other attributes
                if line.startswith('   描述:'):
                    current_event.description = line.replace('   描述:', '').strip()
                elif line.startswith('   地点:'):
                    current_event.location = line.replace('   地点:', '').strip()
                elif line.startswith('   备注:'):
                    current_event.notes = line.replace('   备注:', '').strip()
        
        # Add the last event
        if current_event:
            schedule.add_event(current_event)
        
        return schedule


class WordHandler:
    """Handle Word document import/export."""
    
    @staticmethod
    def export_to_word(schedule: Schedule, file_path: str) -> None:
        """Export schedule to Word document."""
        doc = Document()
        
        # Title
        title = doc.add_heading(f'日程安排: {schedule.name}', 0)
        
        # Schedule info
        if schedule.description:
            doc.add_paragraph(f'描述: {schedule.description}')
        doc.add_paragraph(f'创建时间: {schedule.created_at.strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'更新时间: {schedule.updated_at.strftime("%Y-%m-%d %H:%M:%S")}')
        
        doc.add_paragraph()  # Empty line
        
        # Events
        for i, event in enumerate(schedule.events, 1):
            # Event title
            event_heading = doc.add_heading(f'{i}. {event.title}', level=1)
            
            # Event details table
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            def add_row(label: str, value: str):
                row = table.add_row().cells
                row[0].text = label
                row[1].text = value
            
            add_row('开始时间', event.start_datetime.strftime('%Y-%m-%d %H:%M:%S'))
            if event.end_datetime:
                add_row('结束时间', event.end_datetime.strftime('%Y-%m-%d %H:%M:%S'))
            if event.location:
                add_row('地点', event.location)
            add_row('优先级', event.priority.value)
            add_row('状态', event.status.value)
            if event.tags:
                add_row('标签', ', '.join(event.tags))
            if event.attendees:
                add_row('参与者', ', '.join(event.attendees))
            if event.description:
                add_row('描述', event.description)
            if event.notes:
                add_row('备注', event.notes)
            
            doc.add_paragraph()  # Empty line
        
        doc.save(file_path)
    
    @staticmethod
    def import_from_word(file_path: str) -> Schedule:
        """Import schedule from Word document (basic parsing)."""
        doc = Document(file_path)
        schedule = Schedule(name=f"Imported from {Path(file_path).name}")
        
        current_event = None
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Check if this is an event title (starts with number)
            if text and text[0].isdigit() and '. ' in text:
                if current_event:
                    schedule.add_event(current_event)
                
                title = text.split('. ', 1)[1]
                current_event = ScheduleEvent(
                    title=title,
                    start_datetime=datetime.now()
                )
        
        # Process tables for event details
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) >= 2:
                    label = row.cells[0].text.strip()
                    value = row.cells[1].text.strip()
                    
                    if current_event and value:
                        if label == '开始时间':
                            try:
                                current_event.start_datetime = datetime.strptime(value, '%Y-%m-%d %H:%M:%S')
                            except:
                                pass
                        elif label == '结束时间':
                            try:
                                current_event.end_datetime = datetime.strptime(value, '%Y-%m-%d %H:%M:%S')
                            except:
                                pass
                        elif label == '地点':
                            current_event.location = value
                        elif label == '描述':
                            current_event.description = value
                        elif label == '备注':
                            current_event.notes = value
                        elif label == '标签':
                            current_event.tags = [tag.strip() for tag in value.split(',')]
                        elif label == '参与者':
                            current_event.attendees = [att.strip() for att in value.split(',')]
        
        # Add the last event
        if current_event:
            schedule.add_event(current_event)
        
        return schedule